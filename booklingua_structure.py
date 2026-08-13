"""booklingua_structure.py — Detect headings and apply professional book typography.

Belt-and-braces heading detection:
  1. Source DOCX styles (free, instant)
  2. EPUB markers ###H1:–###H6: (free, instant)
  3. Heuristics + optional Claude Haiku validation (~10¢/book)

Also contains EPUB_STYLESHEET for generate_epub().
"""

import os
import re
import json
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ── Heading Map ────────────────────────────────────────────────────────────

class HeadingMap:
    """Maps paragraph index → heading level (1-3) and detection source."""
    def __init__(self):
        self.levels: Dict[int, int] = {}    # para_idx → 1|2|3
        self.sources: Dict[int, str] = {}   # para_idx → 'source_style'|'source_epub_marker'|'heuristic'|'claude'

    def summary(self) -> str:
        if not self.levels:
            return "0 headings detected"
        h1 = sum(1 for v in self.levels.values() if v == 1)
        h2 = sum(1 for v in self.levels.values() if v == 2)
        h3 = sum(1 for v in self.levels.values() if v == 3)
        src = self.sources.get(next(iter(self.levels)), "unknown")
        return f"{len(self.levels)} headings (H1={h1}, H2={h2}, H3={h3}; source={src})"


# ── Belt 1: Source DOCX styles ─────────────────────────────────────────────

def detect_from_source_styles(source_doc: Document) -> HeadingMap:
    """Trust Heading 1/2/3 styles in the source manuscript."""
    result = HeadingMap()
    for i, para in enumerate(source_doc.paragraphs):
        style = para.style.name.lower()
        if 'heading 1' in style:
            result.levels[i] = 1
            result.sources[i] = 'source_style'
        elif 'heading 2' in style:
            result.levels[i] = 2
            result.sources[i] = 'source_style'
        elif 'heading 3' in style:
            result.levels[i] = 3
            result.sources[i] = 'source_style'
    return result


# ── Belt 2: EPUB markers ───────────────────────────────────────────────────

_EPUB_MARKER = re.compile(r'^###(?:H([123456])|CHAPTER):(.+)###$')


def detect_from_epub_markers(paragraphs: List[str]) -> HeadingMap:
    """
    Detect headings from ###H1:Title### or ###CHAPTER:Title### markers.
    Handles both new H-marker format and legacy CHAPTER format.
    """
    result = HeadingMap()
    for i, text in enumerate(paragraphs):
        m = _EPUB_MARKER.match(text.strip())
        if m:
            level = int(m.group(1)) if m.group(1) else 1
            result.levels[i] = min(level, 3)  # cap at H3
            result.sources[i] = 'source_epub_marker'
    return result


# ── Braces: Heuristic candidate detection ──────────────────────────────────

_HEURISTIC_RE = re.compile(
    r'^'                           # start of line
    r'(?!\d+\.)'                   # NOT a numbered list item
    r'(?!\s*[-•])'                 # NOT a bullet list item
    r'(?!https?://)'               # NOT a URL
    r'(?!.*www\.)'                 # NOT a web address
    r'.{3,180}'                    # 3–180 chars
    r'[^.!?]$'                     # NO terminal punctuation
    r'$',
    re.UNICODE
)


def detect_heading_candidates(paragraphs: List[str]) -> List[int]:
    """
    Return indices of paragraphs that *look* like headings by shape alone.
    Filters out: numbered lists, bullets, URLs, very short lines, sentences.
    """
    candidates = []
    for i, text in enumerate(paragraphs):
        stripped = text.strip()
        if not stripped:
            continue
        # Basic shape filter
        if not _HEURISTIC_RE.match(stripped):
            continue
        # Must have multiple words
        if len(stripped.split()) < 2:
            continue
        # Must not be all lowercase (titles are title case or all caps)
        if stripped == stripped.lower() and stripped != stripped.upper():
            continue
        candidates.append(i)
    return candidates


# ── Claude Haiku validation ────────────────────────────────────────────────

_CLAUDE_PROMPT = """You are a book typesetter. Given a list of paragraph candidates from a translated book, classify each as:
- h1 = chapter title (major division, often starts with "Chapter N" or is a standalone title)
- h2 = section heading (subdivision within a chapter)
- h3 = minor heading (sub-subdivision)
- body = normal body text (NOT a heading)

Context for each candidate: the candidate text plus the first 80 chars of the next paragraph.

Reply with one line per candidate in this exact format:
INDEX:LEVEL
where LEVEL is h1, h2, h3, or body.

Examples:
0:h1
1:body
2:h2
3:body
"""


def validate_with_claude(
    paragraphs: List[str],
    candidates: List[int],
    target_language: str,
    anthropic_client=None,
) -> Dict[int, int]:
    """
    Send candidates in batches of 40 to Claude Haiku for validation.
    Returns {para_idx: level} for confirmed headings (level 1-3).
    """
    if not candidates:
        return {}

    if anthropic_client is None:
        from anthropic import Anthropic
        anthropic_client = Anthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))

    confirmed: Dict[int, int] = {}
    BATCH_SIZE = 40

    for batch_start in range(0, len(candidates), BATCH_SIZE):
        batch = candidates[batch_start:batch_start + BATCH_SIZE]
        lines = []
        for idx in batch:
            text = paragraphs[idx].strip()
            next_text = ""
            if idx + 1 < len(paragraphs):
                next_text = paragraphs[idx + 1].strip()[:80]
            lines.append(f"{idx}: {text} | NEXT: {next_text}")

        prompt = _CLAUDE_PROMPT + "\n\n" + "\n".join(lines)

        try:
            model_config = json.loads((Path(__file__).parent / "config" / "booklingua-models.json").read_text())
            resp = anthropic_client.messages.create(
                model=model_config["normal"],
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}],
            )
            content = resp.content[0].text if resp.content else ""
            for line in content.strip().split("\n"):
                line = line.strip()
                if ":" not in line:
                    continue
                parts = line.split(":", 1)
                if len(parts) != 2:
                    continue
                try:
                    idx = int(parts[0].strip())
                    level_str = parts[1].strip().lower()
                except ValueError:
                    continue
                if level_str == "body":
                    continue
                if level_str in ("h1", "h2", "h3"):
                    confirmed[idx] = int(level_str[1])
        except Exception as e:
            print(f"[Claude validation error] {e}")
            continue

    return confirmed


# ── Orchestrator: detect_headings ──────────────────────────────────────────

def detect_headings(
    source_docx_path: Optional[str],
    translated_paragraphs: List[str],
    target_language: str,
    source_epub_paragraphs: Optional[List[str]] = None,
    use_claude: bool = True,
    anthropic_client=None,
) -> HeadingMap:
    """
    Belt-and-braces heading detection.
    Stops at first belt that finds headings.
    """
    # Belt 1: source DOCX styles
    if source_docx_path and os.path.exists(source_docx_path):
        source_doc = Document(source_docx_path)
        result = detect_from_source_styles(source_doc)
        if result.levels:
            return result

    # Belt 2: EPUB markers (new + legacy formats)
    if source_epub_paragraphs:
        result = detect_from_epub_markers(source_epub_paragraphs)
        if result.levels:
            return result

    # Also try on translated paragraphs (markers may have survived translation)
    result = detect_from_epub_markers(translated_paragraphs)
    if result.levels:
        return result

    # Braces: heuristics + optional Claude validation
    candidates = detect_heading_candidates(translated_paragraphs)
    if not candidates:
        return HeadingMap()

    if not use_claude:
        result = HeadingMap()
        for idx in candidates:
            result.levels[idx] = 2
            result.sources[idx] = 'heuristic'
        return result

    confirmed = validate_with_claude(
        translated_paragraphs, candidates, target_language,
        anthropic_client=anthropic_client,
    )
    result = HeadingMap()
    for idx, level in confirmed.items():
        result.levels[idx] = level
        result.sources[idx] = 'claude'
    return result


# ── Styling: apply to DOCX ─────────────────────────────────────────────────

def _ensure_style(doc: Document, name: str, base_style: str):
    """Get or create a paragraph style."""
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_style]
    return style


def style_translated_book(
    source_docx_path: Optional[str],
    translated_docx_path: str,
    output_docx_path: str,
    target_language: str,
    use_claude: bool = True,
    anthropic_client=None,
    source_epub_paragraphs: Optional[List[str]] = None,
) -> HeadingMap:
    """
    Main entry point:
      1. Detect headings
      2. Apply EB Garamond typography
      3. Save styled DOCX
    """
    # Load translated document
    doc = Document(translated_docx_path)
    paragraphs = [p.text for p in doc.paragraphs]

    # Detect headings
    heading_map = detect_headings(
        source_docx_path=source_docx_path,
        translated_paragraphs=paragraphs,
        target_language=target_language,
        source_epub_paragraphs=source_epub_paragraphs,
        use_claude=use_claude,
        anthropic_client=anthropic_client,
    )

    # Apply styles
    for i, para in enumerate(doc.paragraphs):
        if i in heading_map.levels:
            level = heading_map.levels[i]
            if level == 1:
                para.style = 'Heading 1'
                para.paragraph_format.space_before = Pt(18)
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.page_break_before = True
                para.paragraph_format.keep_with_next = True
                for run in para.runs:
                    run.font.name = 'EB Garamond'
                    run.font.size = Pt(22)
                    run.font.bold = True
            elif level == 2:
                para.style = 'Heading 2'
                para.paragraph_format.space_before = Pt(18)
                para.paragraph_format.space_after = Pt(10)
                para.paragraph_format.keep_with_next = True
                for run in para.runs:
                    run.font.name = 'EB Garamond'
                    run.font.size = Pt(16)
                    run.font.bold = True
            elif level == 3:
                para.style = 'Heading 3'
                para.paragraph_format.space_before = Pt(14)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.keep_with_next = True
                for run in para.runs:
                    run.font.name = 'EB Garamond'
                    run.font.size = Pt(13)
                    run.font.bold = True
        else:
            # Body text
            para.paragraph_format.line_spacing = 1.4
            para.paragraph_format.first_line_indent = Inches(0.3)
            for run in para.runs:
                run.font.name = 'EB Garamond'
                run.font.size = Pt(11)

        # First paragraph after heading: no first-line indent
        if i > 0 and (i - 1) in heading_map.levels:
            para.paragraph_format.first_line_indent = Inches(0)

    doc.save(output_docx_path)
    return heading_map


# ── EPUB Stylesheet ────────────────────────────────────────────────────────

EPUB_STYLESHEET = """/* content.css — BookLingua EPUB stylesheet */

@font-face {
  font-family: 'EB Garamond';
  src: url('../fonts/EBGaramond-Regular.ttf') format('truetype');
  font-weight: normal;
  font-style: normal;
}

@font-face {
  font-family: 'EB Garamond';
  src: url('../fonts/EBGaramond-Bold.ttf') format('truetype');
  font-weight: bold;
  font-style: normal;
}

body {
  font-family: 'EB Garamond', Garamond, Georgia, serif;
  font-size: 11pt;
  line-height: 1.4;
  text-align: justify;
  margin: 0;
  padding: 0;
}

p {
  text-indent: 1.5em;
  margin: 0;
  orphans: 2;
  widows: 2;
}

h1 {
  font-family: 'EB Garamond', Garamond, Georgia, serif;
  font-size: 22pt;
  font-weight: bold;
  text-align: center;
  margin-top: 2em;
  margin-bottom: 1em;
  page-break-before: always;
  page-break-after: avoid;
  text-indent: 0;
}

h2 {
  font-family: 'EB Garamond', Garamond, Georgia, serif;
  font-size: 16pt;
  font-weight: bold;
  margin-top: 1.5em;
  margin-bottom: 0.6em;
  page-break-after: avoid;
  text-indent: 0;
}

h3 {
  font-family: 'EB Garamond', Garamond, Georgia, serif;
  font-size: 13pt;
  font-weight: bold;
  margin-top: 1.2em;
  margin-bottom: 0.4em;
  page-break-after: avoid;
  text-indent: 0;
}

/* First paragraph after heading: no indent */
h1 + p, h2 + p, h3 + p {
  text-indent: 0;
}

/* Title page */
.title {
  font-size: 28pt;
  font-weight: bold;
  text-align: center;
  margin-top: 30%;
  text-indent: 0;
}

.subtitle {
  font-size: 14pt;
  text-align: center;
  margin-top: 1em;
  text-indent: 0;
}

.author {
  font-size: 13pt;
  text-align: center;
  margin-top: 2em;
  text-indent: 0;
}

/* Copyright page */
.copyright {
  font-size: 9pt;
  text-align: center;
  margin-top: 40%;
  text-indent: 0;
}
"""
