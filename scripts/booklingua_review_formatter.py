"""
booklingua_review_formatter.py

Post-processing pass on review DOCXs before they are sent for human review.

Takes a raw review DOCX produced by the pipeline and:
  1. Pulls all ===TRANSLATION_NOTES=== blocks to a summary section at the top
  2. Converts [[ORIGINAL: old]]new markers to yellow-highlighted tracked changes
     (strikethrough original in yellow, arrow, new text in yellow)
  3. Strips the inline ===TRANSLATION_NOTES=== markers from the body
  4. Adds a clean header with change count and date

Usage:
    python booklingua_review_formatter.py input.docx output.docx [--lang de]

Or import and call format_review_docx() directly from the pipeline.
"""

from __future__ import annotations

import argparse
import copy
import os
import re
import sys
from datetime import date
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX


# ── Constants ─────────────────────────────────────────────────────────────

ORIGINAL_RE = re.compile(r'\[\[ORIGINAL:\s*(.*?)\]\]', re.DOTALL)
TN_MARKER_RE = re.compile(r'^===TRANSLATION_NOTES===\s*$')
SEGMENT_RE = re.compile(r'^===SEGMENT_\d+_(START|END)===\s*$')
SEGMENT_INLINE_RE = re.compile(r'===SEGMENT_\d+_(START|END)===')
YELLOW = "yellow"
BRAND = "1F3864"
ACCENT = "C0392B"
SECTION_COLOR = "2E5D8E"


# ── XML run builder ───────────────────────────────────────────────────────

def _make_run(text: str, bold: bool = False, italic: bool = False,
              highlight: Optional[str] = None,
              strikethrough: bool = False,
              color: Optional[str] = None,
              size_pt: Optional[int] = None) -> OxmlElement:
    """Build a w:r element with text and optional formatting."""
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    if bold:
        rpr.append(OxmlElement('w:b'))
    if italic:
        rpr.append(OxmlElement('w:i'))
    if strikethrough:
        rpr.append(OxmlElement('w:strike'))
    if highlight:
        h = OxmlElement('w:highlight')
        h.set(qn('w:val'), highlight)
        rpr.append(h)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rpr.append(c)
    if size_pt:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_pt * 2))
        rpr.append(sz)
        sz2 = OxmlElement('w:szCs')
        sz2.set(qn('w:val'), str(size_pt * 2))
        rpr.append(sz2)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r


def _add_highlighted_runs_to_para(para, text: str) -> int:
    """Add text with [[ORIGINAL:]] markers to a paragraph as highlighted runs.
    Uses docx API (not OxmlElement). Returns number of changes found."""
    if '[[ORIGINAL:' not in text:
        para.add_run(text)
        return 0

    all_matches = list(ORIGINAL_RE.finditer(text))
    if not all_matches:
        para.add_run(text)
        return 0

    pos = 0
    count = 0
    for idx, m in enumerate(all_matches):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        orig_text = m.group(1).strip()
        next_start = all_matches[idx + 1].start() if idx + 1 < len(all_matches) else len(text)
        replacement = text[m.end():next_start]

        if orig_text:
            r = para.add_run(orig_text)
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            r.font.strike = True
            r = para.add_run(' → ')
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if replacement:
            r = para.add_run(replacement)
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        pos = next_start
        count += 1

    if pos < len(text):
        para.add_run(text[pos:])

    return count


# ── Paragraph highlighter (legacy, for body paragraphs) ───────────────────

def rewrite_paragraph_with_highlights(para) -> int:
    """
    Convert [[ORIGINAL: old]]new markers in a paragraph to yellow highlights.
    Returns number of changes made.
    """
    full_text = para.text
    if '[[ORIGINAL:' not in full_text:
        return 0

    all_matches = list(ORIGINAL_RE.finditer(full_text))
    if not all_matches:
        return 0

    # Build token list: ('plain', text) or ('change', orig, replacement)
    tokens = []
    pos = 0
    for idx, m in enumerate(all_matches):
        if m.start() > pos:
            tokens.append(('plain', full_text[pos:m.start()]))
        orig_text = m.group(1).strip()
        next_start = all_matches[idx + 1].start() if idx + 1 < len(all_matches) else len(full_text)
        replacement = full_text[m.end():next_start]
        tokens.append(('change', orig_text, replacement))
        pos = next_start

    # Remove all existing runs from the paragraph
    p_elem = para._p
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink', 'ins', 'del'):
            p_elem.remove(child)

    # Add new runs
    for token in tokens:
        if token[0] == 'plain' and token[1]:
            p_elem.append(_make_run(token[1]))
        elif token[0] == 'change':
            orig_text, replacement = token[1], token[2]
            if orig_text:
                p_elem.append(_make_run(orig_text, highlight=YELLOW, strikethrough=True))
                p_elem.append(_make_run(' → ', highlight=YELLOW))
            if replacement:
                p_elem.append(_make_run(replacement, highlight=YELLOW))

    return len(all_matches)


# ── Inline segment marker stripper ──────────────────────────────────────

def strip_inline_segment_markers(doc: Document) -> int:
    """Remove inline ===SEGMENT_*=== markers from all paragraph text."""
    count = 0
    for para in doc.paragraphs:
        full_text = para.text
        if '===SEGMENT_' not in full_text:
            continue
        cleaned = SEGMENT_INLINE_RE.sub('', full_text)
        if cleaned != full_text:
            p_elem = para._p
            bold = False
            italic = False
            for run in para.runs:
                rpr = run._element.find(qn('w:rPr'))
                if rpr is not None:
                    if rpr.find(qn('w:b')) is not None:
                        bold = True
                    if rpr.find(qn('w:i')) is not None:
                        italic = True
            for child in list(p_elem):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('r', 'hyperlink', 'ins', 'del'):
                    p_elem.remove(child)
            p_elem.append(_make_run(cleaned, bold=bold, italic=italic))
            count += 1
    return count


# ── Translation notes extractor ───────────────────────────────────────────

def extract_translation_notes(doc: Document) -> tuple[list[str], set[int]]:
    """
    Extract all ===TRANSLATION_NOTES=== blocks.
    Returns (notes_blocks, set_of_paragraph_indices_to_remove).
    """
    paragraphs = doc.paragraphs
    notes_blocks = []
    to_remove = set()

    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if TN_MARKER_RE.match(p.text.strip()):
            to_remove.add(i)
            i += 1
            block_lines = []
            while i < len(paragraphs):
                np = paragraphs[i]
                if TN_MARKER_RE.match(np.text.strip()):
                    break
                try:
                    if np.style.name in ('Heading 1', 'Heading 2') and np.text.strip():
                        break
                except Exception:
                    pass
                to_remove.add(i)
                if np.text.strip():
                    block_lines.append(np.text)
                i += 1
            if block_lines:
                notes_blocks.append('\n'.join(block_lines))
        else:
            i += 1

    return notes_blocks, to_remove


# ── Header builder (uses doc API for reliable text rendering) ─────────────

def prepend_header(doc: Document, title: str, language: str,
                   change_count: int, notes_blocks: list[str]) -> None:
    """
    Add a formatted header section at the very top of the document.
    Uses doc.add_paragraph() for reliable text rendering, then moves
    the added paragraphs to the front of the document body.
    """
    body = doc.element.body
    # Record the current first child position
    first_existing = list(body)[0]

    added = []

    def ap(text: str = '', bold: bool = False, italic: bool = False,
           color: Optional[str] = None, size_pt: Optional[int] = None) -> None:
        """Add a paragraph via the doc API, track it for moving."""
        p = doc.add_paragraph()
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
            if size_pt:
                run.font.size = Pt(size_pt)
        added.append(p._p)

    def divider() -> None:
        p = doc.add_paragraph()
        ppr = OxmlElement('w:pPr')
        pbdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'AAAAAA')
        pbdr.append(bot)
        ppr.append(pbdr)
        p._p.insert(0, ppr)
        added.append(p._p)

    def highlight_legend() -> None:
        """Add a paragraph with the yellow highlight colour key."""
        p = doc.add_paragraph()
        p._p.append(_make_run('■  ', highlight=YELLOW, bold=True, size_pt=11))
        p._p.append(_make_run(
            'Yellow = translator made a change   |   ',
            size_pt=11
        ))
        p._p.append(_make_run('Strikethrough', highlight=YELLOW,
                               strikethrough=True, size_pt=11))
        p._p.append(_make_run(' = original text   |   Plain yellow = new text', size_pt=11))
        added.append(p._p)

    # ── Header content ──
    ap()
    ap(f"BookLingua Review Document", bold=True, size_pt=18, color=BRAND)
    ap(f"{title}  ·  {language}", bold=True, size_pt=14)
    ap(f"Generated: {date.today().strftime('%d %B %Y')}")
    ap(
        f"{change_count} translation changes flagged   |   "
        f"{len(notes_blocks)} translator note blocks",
        bold=True, color=ACCENT
    )
    ap()
    ap("COLOUR KEY", bold=True, color="555555")
    highlight_legend()
    ap()
    divider()
    ap()

    if notes_blocks:
        ap("TRANSLATION NOTES SUMMARY", bold=True, size_pt=14, color=BRAND)
        ap(
            "All translator commentary consolidated below. "
            "Removed from the body of the document.",
            italic=True, color="555555"
        )
        ap()
        for chapter_num, block in enumerate(notes_blocks, 1):
            ap(f"─── Chapter {chapter_num} Notes ───", bold=True, color=SECTION_COLOR)
            for line in block.split('\n'):
                if line.strip():
                    # Highlight [[ORIGINAL:]] markers in notes too
                    if '[[ORIGINAL:' in line:
                        p = doc.add_paragraph()
                        _add_highlighted_runs_to_para(p, line.strip())
                        added.append(p._p)
                    else:
                        ap(line.strip())
            ap()

    divider()
    ap()
    ap("TRANSLATION TEXT", bold=True, size_pt=14, color=BRAND)
    ap(
        "Changes highlighted below. "
        "Strikethrough = original text  ·  Plain yellow = translator's version."
    )
    ap()

    # Move all added paragraphs to the front of the body
    for p_elem in reversed(added):
        body.remove(p_elem)
        body.insert(0, p_elem)


# ── Main formatter ────────────────────────────────────────────────────────

def format_review_docx(input_path: str, output_path: str,
                        language: str = "",
                        title: str = "") -> dict:
    """
    Format a raw review DOCX for human review.

    Args:
        input_path:  path to the raw review DOCX from the pipeline
        output_path: path to write the formatted output
        language:    human-readable language name e.g. "German"
        title:       book title

    Returns:
        dict with stats: change_count, note_blocks, paragraphs_processed
    """
    doc = Document(input_path)

    # Infer title and language
    if not title and doc.paragraphs:
        title = doc.paragraphs[0].text.strip() or "Unknown"
    if not language:
        fn = os.path.basename(input_path).lower()
        if 'german' in fn or '_de' in fn:   language = "German"
        elif 'italian' in fn or '_it' in fn: language = "Italian"
        elif 'spanish' in fn or '_es' in fn: language = "Spanish (Latin America)"
        elif 'french' in fn or '_fr' in fn:  language = "French"
        else:                                 language = "Translation"

    # Step 0: Strip inline segment markers from all paragraphs
    print(f"  Stripping inline segment markers...")
    stripped = strip_inline_segment_markers(doc)
    print(f"  Stripped {stripped} paragraphs with inline segment markers")

    # Step 1: Extract translation notes
    print(f"  Extracting translation notes...")
    notes_blocks, tn_indices = extract_translation_notes(doc)
    print(f"  Found {len(notes_blocks)} note blocks ({len(tn_indices)} paragraphs)")

    # Step 1b: Count [[ORIGINAL:]] markers inside notes blocks (non-fiction pattern)
    notes_changes = 0
    for block in notes_blocks:
        notes_changes += len(ORIGINAL_RE.findall(block))
    print(f"  Found {notes_changes} changes inside notes blocks")

    # Step 2: Highlight [[ORIGINAL:]] markers in body paragraphs
    # Only process paragraphs NOT in the translation notes blocks
    print(f"  Highlighting translation changes...")
    total_changes = 0
    segment_indices = set()
    for i, para in enumerate(doc.paragraphs):
        if SEGMENT_RE.match(para.text.strip()):
            segment_indices.add(i)
            continue
        if i in tn_indices:
            continue
        if TN_MARKER_RE.match(para.text.strip()):
            continue
        n = rewrite_paragraph_with_highlights(para)
        total_changes += n
    total_changes += notes_changes
    print(f"  Total changes (body + notes): {total_changes}")

    # Step 3: Remove TN paragraphs + segment markers from body (reverse order)
    print(f"  Removing inline translation note blocks + segment markers...")
    body = doc.element.body
    paragraphs = doc.paragraphs
    removed = 0
    for idx in sorted(tn_indices | segment_indices, reverse=True):
        if idx < len(paragraphs):
            p_elem = paragraphs[idx]._p
            parent = p_elem.getparent()
            if parent is not None:
                parent.remove(p_elem)
                removed += 1
    print(f"  Removed {removed} paragraphs ({len(segment_indices)} segment markers)")

    # Step 4: Prepend the header section
    print(f"  Building header section...")
    prepend_header(
        doc=doc,
        title=title,
        language=language,
        change_count=total_changes,
        notes_blocks=notes_blocks,
    )

    # Step 5: Save
    doc.save(output_path)
    print(f"  Saved to {output_path}")

    return {
        "change_count": total_changes,
        "body_changes": total_changes - notes_changes,
        "notes_changes": notes_changes,
        "note_blocks": len(notes_blocks),
        "tn_paragraphs_removed": removed,
        "segment_markers_removed": len(segment_indices) + stripped,
        "language": language,
        "title": title,
    }


# ── CLI ───────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Format a BookLingua review DOCX for human review"
    )
    parser.add_argument("input", help="Input review DOCX path")
    parser.add_argument("output", help="Output formatted DOCX path")
    parser.add_argument("--lang", default="", help="Language name e.g. 'German'")
    parser.add_argument("--title", default="", help="Book title")
    args = parser.parse_args()

    print(f"Formatting review DOCX: {args.input}")
    stats = format_review_docx(
        input_path=args.input,
        output_path=args.output,
        language=args.lang,
        title=args.title,
    )
    print(f"\nDone.")
    print(f"  Language:         {stats['language']}")
    print(f"  Changes flagged:  {stats['change_count']}")
    print(f"  Note blocks:      {stats['note_blocks']}")
    print(f"  TN paras removed: {stats['tn_paragraphs_removed']}")
    print(f"  Output:           {args.output}")


if __name__ == "__main__":
    main()
